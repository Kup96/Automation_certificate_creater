import io
 
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage

from docx import Document
from docx.shared import Inches

import re
import datetime

print ('Нажмите любую кнопку для начала работы')
input ('')
 
def extract_text_from_pdf(pdf_path):
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle)
    page_interpreter = PDFPageInterpreter(resource_manager, converter)


    with open(pdf_path, 'rb') as fh:
        for page in PDFPage.get_pages(fh, 
                                      caching=True,
                                      check_extractable=True):
            page_interpreter.process_page(page)
 
        text = fake_file_handle.getvalue()
 
    # close open handles
    converter.close()
    fake_file_handle.close()
 
    if text:
        return text
 
if __name__ == '__main__':
    
    score_without_head = extract_text_from_pdf('turbo.pdf')
score_without_ass = score_without_head[score_without_head.find("без ПДВ") + 18 : ]
index_of_last = (score_without_ass.rfind("шт")+2)
score_without_ass_and_head = score_without_ass[:index_of_last]
seals = re.sub(r'шток', "торк", score_without_ass_and_head)
score_str = re.sub(r'(?<=(шт))(.+?)(?=\")', '\n', seals)
score_str_seals = re.sub(r'торк', "шток", score_str)
score_str_sqmetres = re.sub (r'(?<=(кв\.м))(.+?)(?=\")', '\n',score_str_seals)
score_str_kg = re.sub(r'(?<=(кг))(.+?)(?=\")', '\n', score_str_sqmetres)
score_str_pog3 = re.sub(r'(?<=(пог\.м))(.+?)(?=\")', '\n', score_str_kg)
score_str_pog = re.sub(r'(?<=(м))(.+?)(?=\")', '\n', score_str_pog3)
len_of_future_tables = len(re.findall(r"[\n']+?", score_str_pog))
index_of_lasts = (score_str_kg.rfind("шт" and "кг")+2)
number_of_score = re.search (r'\d{1,4}', score_without_head)


def set_column_width(table,column,width_inch):
    table.allow_autofit = False
    for row in table.rows:
        row.cells[column].width = Inches(width_inch)

k = 'К?льце'
m = 'Манжета'
b = 'Брудоз\'ємник'
g = 'Грязес\'ємник'
c = 'Шнур'


def check(score_str_kg):
	vivod = ""
	if re.search(f'{k}', score_str_kg):
		vivod += str("Кільце")
	if re.search(f'{m}', score_str_kg):
		vivod +=str(" Манжета")
	if re.search(f'{b}', score_str_kg):
		vivod +=str(' Брудоз\'ємник')
	if re.search(f'{g}', score_str_kg):
		vivod +=str('Грязес\'ємник')
	return vivod

product_name = check(score_str_kg)

date_today = datetime.datetime.today().strftime("%d%m")
date_today_j = datetime.datetime.today().strftime("%d")
date_yesterday = datetime.datetime.today() - datetime.timedelta(days=1)
date_yesterday =  date_yesterday.strftime("%d")
month_today = (datetime.datetime.today().strftime("%m"))
partiya_date = date_yesterday + month_today

document = Document()
document.add_picture('shapka_ukr.jpg', width=Inches(7.086))
sertificat = document.add_paragraph('                                                             ')
sertificat.add_run('СЕРТИФІКАТ ЯКОСТІ №').bold = True
sertificat.add_run(str(date_today)).bold = True
name_of_product = document.add_paragraph ("")
name_of_product.add_run ("Наименування партії: ").bold = True
name_of_product.add_run (product_name).underline = True
partiya = document.add_paragraph('')
partiya.add_run ("Партія № " + partiya_date).bold = True
counts_kg = re.findall(r'\d{1,4}кг',score_str_kg)
counts_sqm = re.findall(r'\d{1,4}кв\.м',score_str_kg)
counts1 = re.findall(r'\d{1,4}м', score_str_kg)
counts = re.findall(r'\d{1,4}шт', score_str_kg)
score_str_kg_new = re.sub(r'\d{1,4}кг', ' ', score_str_kg)
score_str_sqm4 = re.sub(r'\d{1,4}кв\.м', ' ', score_str_kg_new)
score_str_sqm = re.sub(r'\d{1,4}м', ' ', score_str_sqm4)
score_str_kg_normal = re.sub(r'\d{1,4}шт', ' ', score_str_sqm)
counts_kg_new = counts_kg + counts_sqm + counts1 + counts
mists = score_str_kg_normal.split('\n')
counts_kg_new_eng = [re.sub('шт', "" , count) for count in counts_kg_new ]


def do_list_for_birks(mist):
	mist = re.sub(r'/ поліамид 6 блочний', '', mist)
	mist = re.sub(r'/поліамид 6 блочний', '', mist)
	mist = re.sub(r'поліамид 6 блочний', '', mist)
	mist = re.sub(r'/ГОСТ14896-84', '', mist)
	mist = re.sub(r'/ ГОСТ14896-84', '', mist)
	mist = re.sub(r'ГОСТ14896-84', '', mist)
	mist = re.sub(r'ГОСТ 14896-84', '', mist)
	mist = re.sub(r'/ГОСТ18829-73', '', mist)
	mist = re.sub(r'/ ГОСТ18829-73', '', mist)
	mist = re.sub(r'ГОСТ18829-73', '', mist)
	mist = re.sub(r'ГОСТ 18829-73', '', mist)
	mist = re.sub(r'ГОСТ22704-77', '', mist)
	mist = re.sub(r'/ГОСТ22704-77', '', mist)
	mist = re.sub(r'/ ГОСТ22704-77', '', mist)
	mist = re.sub(r'ГОСТ 22704-77', '', mist)
	mist = re.sub(r'/ ГОСТ 9864-86', '', mist)
	mist = re.sub(r'/ГОСТ 9864-86', '', mist)
	mist = re.sub(r'ГОСТ 9864-86', '', mist)
	mist = re.sub(r'ГОСТ9864-86', '', mist)
	mist = re.sub('/ Форполімер уретановий', '', mist)
	mist = re.sub('/ форполімер уретановий', '', mist)
	mist = re.sub ('форполімер уретановий', '', mist)
	mist = re.sub('Форполімер уретановий', '', mist)
	mist = re.sub('/рез. гр. III', '', mist)
	mist = re.sub('/ рез. гр. III', '', mist)
	mist = re.sub('рез. гр. III', '', mist)
	return mist

mist_birks = list(map(do_list_for_birks, mists))

def do_list_material(mist):
	if re.search ('Кільце' and "опорне" or "КО", mist):
		mist = "Поліамід 6 блочний "
	elif re.search ('Кільце' and "нажимне" or "КН", mist):
		mist = "Полiамід 6 блочний "
	elif re.search ('Манжета' and '198964', mist):
		mist = "Форполiмер уретановий"
	elif re.search ('Манжета' and '197787', mist):
		mist = "Форполiмер уретановий"
	elif re.search ("форполімер уретановий", mist):
		mist = "Форполімер уретановий"
	elif re.search ("Поліамід 6 блочний", mist):
		mist = "Поліамід 6 блочний"
	elif re.search ('Кільце' and '2-6', mist):
		mist = 'Силiконовий каучук'
	elif re.search ('Кільце', mist):
		mist = "7-В-14"
	elif re.search ('3-ГТ200598', mist):
		mist = "Резина гр. VI, марка 7-В-14-1"
	elif re.search ('Манжета', mist):
		mist = "Резина гр. III, марка 7-В-14-1"
	else:
		mist = "  "
	return mist


def do_list_material_eng(mist):
	if re.search ('Кільце' and "опорне" or "КО", mist):
		mist = "Polyamide 6A 6C A112"
	elif re.search ('Кільце' and "нажимне" or "КН", mist):
		mist = "Polyamide 6A 6C A112 "
	elif re.search ('Манжета' and '198964', mist):
		mist = "Urethane polyamer"
	elif re.search ('Манжета' and '197787', mist):
		mist = "Urethane polyamer"
	elif re.search ("форполімер уретановий", mist):
		mist = "Urethane polyamer"
	elif re.search ("Поліамід 6 блочний", mist):
		mist = "Polyamide 6A 6C A112"
	elif re.search ('Кільце' and '2-6', mist):
		mist = 'Silicone'
	elif re.search ('Кільце', mist):
		mist = "NBR"
	elif re.search ('Манжета', mist):
		mist = "NBR"
	else:
		mist = "  "
	return mist

def do_list_gost(mist):
	if re.search ('14896-84', mist):
		mist = "14896-84"
	elif re.search ('18829-73', mist):
		mist = "18829-73"
	elif re.search ('М-424-26-12', mist):
		mist = "креслення М-424-26-12"
	elif re.search ('22704-77' or '197787', mist):
		mist = "22704-77"
	elif re.search ('9864-86', mist):
		mist = "9864-86"
	else:
		mist = "  "
	return mist

material =  list(map(do_list_material, mists))
material_eng = list(map(do_list_material_eng, mists))
gostins = list(map(do_list_gost,mists))

passport_table = document.add_table (rows = len_of_future_tables+2, cols = 5)
passport_table.style = 'Table Grid'
set_column_width(passport_table, 0, 0.27)
set_column_width(passport_table, 1, 4.5)
set_column_width(passport_table, 2, 2)
set_column_width(passport_table, 3, 1.5)
set_column_width(passport_table, 4, 0.3)
hdr_cells = passport_table.rows[0].cells
hdr_cells[0].text = '№'
hdr_cells[1].text = 'Найменування'
hdr_cells[2].text = 'Марка\nгуми'
hdr_cells[3].text = '№ партії'
hdr_cells[4].text = 'Кількість\nштук'

for row in range(1,len_of_future_tables+2):
    for col in range(1):
    	cell = passport_table.cell(row, col)
    cell.text = str(row)

for row in range(1,len_of_future_tables+2):
	if row == 10:
		break
	for col in range(4):
    		cell = passport_table.cell(row, col)
	cell.text = str(partiya_date) + '-00' + str(row)
if len_of_future_tables > 9:
	for row in range(10,len_of_future_tables+2):
		for col in range(4):
				cell = passport_table.cell(row, col)
		cell.text = str(partiya_date) + '-0' + str(row)


for row, mist in zip(range(1, len_of_future_tables+2), mists):
    for col in range(2):
        cell = passport_table.cell(row, col)
    cell.text = (mist)

for row, count in zip(range(1, len_of_future_tables+2), counts_kg_new):
    for col in range(5):
        cell = passport_table.cell(row, col)
    cell.text = (count)
for row, mater in zip(range(1, len_of_future_tables+2), material):
    for col in range(3):
        cell = passport_table.cell(row, col)
    cell.text = (mater)

rezultat_of_control = document.add_paragraph('\n                                                             ')
rezultat_of_control.add_run('1.  Результат контролю').bold = True

gost14896 = "14896-84"
gost18829 = "18829-73"
gost22704 = "22704-77"
PU = 'форполімер уретановий'
m4 = 'М-424-26-12'


def gosts(score_str_kg):
	nomera = "ГОСТ "
	if re.search(f'{gost14896}', score_str_kg):
		nomera += str("14896-84")
	if re.search(f'{gost18829}', score_str_kg):
		nomera +=str(" 18829-73")
	if re.search(f'{gost22704}', score_str_kg):
		nomera +=str(' 22704-77')
	if re.search(f'{PU}', score_str_kg):
		nomera +=str(' ТУ 38.105376-92')
	if re.search(f'{m4}', score_str_kg):
		nomera +=str('креслення ТУ 38.105376-92')
	return nomera
def gosts_eng(score_str_kg):
	nomera = "GOST "
	if re.search(f'{gost14896}', score_str_kg):
		nomera += str("14896-84")
	if re.search(f'{gost18829}', score_str_kg):
		nomera +=str(" 18829-73")
	if re.search(f'{gost22704}', score_str_kg):
		nomera +=str(' 22704-77')
	if re.search(f'{PU}', score_str_kg):
		nomera +=str(' TU 38.105376-92')
	if re.search(f'{m4}', score_str_kg):
		nomera +=str('draw М-424-26-12')
	return nomera

fizikal_mech_gosts = gosts(score_str_kg)
fizikal_mech_gosts_eng = gosts_eng(score_str_kg)
fizikal_mech_char = document.add_paragraph('')
fizikal_mech_char.add_run('1.1 Фізико-механічні показники матеріалу відповідають: \n' + fizikal_mech_gosts + "\nРозміри та зовнішній вигляд відповідають вимогам" + " " + fizikal_mech_gosts).bold = True
enough1 = document.add_paragraph('\n                                              ') 
enough1.add_run('Висновок відділу технічного контролю').bold = True
enough2 = document.add_paragraph('\n        ')
enough2.add_run(' Партія № ' + str(partiya_date) + " відповідають вимогам ").underline = True
enough2.add_run(fizikal_mech_gosts).bold = True
months = ['січня', 'лютого', 'березня', 'квітня', 'травня', 'червня', 'липня', 'серпня', 'вересня', 'жовтня', 'листопада', 'грудня']
mesyac = months[(int(month_today)-1) % 12]
enough3 = document.add_paragraph ("")
enough3.add_run ('\nГарантійний термін обічислюється з ' + mesyac + " 20" + datetime.datetime.today().strftime("%y") + 'г' + '\nТермін зберігання діє 5 років' ).italic = True
enough4 = document.add_paragraph ('\nНачальник ОТК                                                                                                                 О.В. Кузнєцов')

document.add_page_break()
document.save('Паспорт на украинском для счета номер ' + number_of_score.group(0) + '.docx')

print ("Паспорт на украинском готов, делаем бирки")

documentqqq = Document()
nn = (len_of_future_tables+1) * 9
birki_tableq = documentqqq.add_table (rows = nn, cols = 2)
birki_tableq.style = 'Table Grid'
set_column_width(birki_tableq, 0, 1.28)
birks = (("", "Найменування", "ГОСТ", "Матеріал", "Кількість, шт", "Дата \nвиробництва", "№ партії", "Печать ОТК \n ", "Дата") * (len_of_future_tables + 1))
for row, birk in zip(range(nn), birks):
    	for col in range(1):
        	cell = birki_tableq.cell(row, col)
    	cell.text = (birk)
for row, count in zip(range(4,nn,9), counts_kg_new):
		for col in range(2):
			cell = birki_tableq.cell(row,col)
		cell.text = (count)
for row, mist in zip(range(1,nn,9), mist_birks):
		for col in range(2):
			cell = birki_tableq.cell(row,col)
		cell.text = (mist)
for row in range(8,nn,9):
		for col in range(2):
			cell = birki_tableq.cell(row,col)
		cell.text = (datetime.datetime.today().strftime("%d.%m.20%y"))
for row in range(5,nn,9):
		for col in range(2):
			cell = birki_tableq.cell(row,col)
		cell.text = (date_yesterday + datetime.datetime.today().strftime(".%m.20%y"))

xt = (len_of_future_tables+1)
j = 0
formerge = 0
while j < xt:
	ab = birki_tableq.cell(formerge, 0)
	ba = birki_tableq.cell(formerge, 1)
	A = ab.merge(ba)
	B = A.add_paragraph('                                     ')
	HH = B.add_run()
	HH.add_picture('shapka2.jpg')
	formerge = formerge + 9
	j = j + 1

partiya_date_stripes = range(1,len_of_future_tables+2)

for row, partiya_date_stripe in zip(range(6,nn,9), partiya_date_stripes):
		for col in range(2):
			cell = birki_tableq.cell(row,col)
		cell.text = (partiya_date + '-0' + str(partiya_date_stripe))
for row, mater in zip(range(3, nn, 9), material):
    	for col in range(2):
        	cell = birki_tableq.cell(row, col)
    	cell.text = (mater)
for row, gost in zip(range(2, nn, 9), gostins):
    	for col in range(2):
        	cell = birki_tableq.cell(row, col)
    	cell.text = (gost)

documentqqq.save('Бирки для счета ' + number_of_score.group(0) + '.docx')

print ('Бирки готовы, делаем паспорт на английском')
score_str_sqmetres_eng1 = re.sub ('Кільце нажимне', 'Pressure ring', score_str_kg_normal)
score_str_sqmetres_eng2 = re.sub ('Кільце опорне', 'Support ring', score_str_sqmetres_eng1)
score_str_sqmetres_eng3 = re.sub ('Захисне кільце', 'Back-up ring', score_str_sqmetres_eng2)
score_str_sqmetres_eng4 = re.sub ('Кільце', 'Ring', score_str_sqmetres_eng3)
score_str_sqmetres_eng5 = re.sub ('Манжета', 'Cuff', score_str_sqmetres_eng4)
score_str_sqmetres_eng6 = re.sub ('ГОСТ', 'GOST', score_str_sqmetres_eng5)
score_str_sqmetres_eng7 = re.sub ('4-ГТ', '4-GT', score_str_sqmetres_eng6)
score_str_sqmetres_eng8 = re.sub ('форполімер уретановий', 'Urethane polyamer', score_str_sqmetres_eng7)
score_str_sqmetres_eng9 = re.sub ('поліамид 6 блочний', 'Polyamide 6A 6C A112', score_str_sqmetres_eng8)
score_str_sqmetres_eng10 = re.sub ('рез. гр.', 'NBR', score_str_sqmetres_eng9)
score_str_sqmetres_eng11 = re.sub ('3-ГТ', '3-GT', score_str_sqmetres_eng10)
score_str_sqmetres_eng13 = re.sub ('Брудоз\'ємник' or 'Брудос\'ємник' or 'Грязес\'ємник', 'Wiper', score_str_sqmetres_eng11)
score_str_sqmetres_eng14 = re.sub ('Шнур', 'Cord', score_str_sqmetres_eng13)
score_str_sqmetres_eng15 = re.sub ('кресл', 'draw', score_str_sqmetres_eng14)
score_str_sqmetres_eng16 = re.sub ('креслення', 'drawing', score_str_sqmetres_eng15)
score_str_sqmetres_eng17 = re.sub ('ГТ', 'GT', score_str_sqmetres_eng16)
score_str_sqmetres_eng18 = re.sub ('Гума', 'NBR', score_str_sqmetres_eng17)
score_str_sqmetres_eng19 = re.sub ('Резина', 'NBR', score_str_sqmetres_eng18)
score_str_sqmetres_eng20 = re.sub ('Силікон' or 'Силикон' or 'силіконовий', 'Silicone', score_str_sqmetres_eng19)
score_str_sqmetres_eng22 = re.sub ('Профиль' or 'Профіль', 'Profile', score_str_sqmetres_eng20)
score_str_sqmetres_eng24 = re.sub ('гідравлічна', '', score_str_sqmetres_eng22)
score_str_sqmetres_eng25 = re.sub ('СДУ Р', 'NBR', score_str_sqmetres_eng24)
score_str_sqmetres_eng26 = re.sub ('Ущільнення штоку', 'Rod seal', score_str_sqmetres_eng25)
score_str_sqmetres_eng27 = re.sub ('Ущільнення поршню', 'Piston seal', score_str_sqmetres_eng26)
score_str_sqmetres_eng28 = re.sub ('Ущільнення циліндра', 'Rotary seal', score_str_sqmetres_eng27)
score_str_sqmetres_eng29 = re.sub ('пневматична', '', score_str_sqmetres_eng28)
score_str_sqmetres_eng30 = re.sub ('ущільнююче', '', score_str_sqmetres_eng29)
score_str_sqmetres_eng31 = re.sub ('Направляюче кільце', 'Guide ring', score_str_sqmetres_eng30)
score_str_sqmetres_eng32 = re.sub ('Ескіз' or 'ескіз', 'draw', score_str_sqmetres_eng31)
score_str_sqmetres_eng34 = re.sub ('поліамід 6 блочний', 'Polyamide 6A 6C A112', score_str_sqmetres_eng32)
score_str_sqmetres_eng35 = re.sub ('Лист гумовий', 'Sheet of NBR', score_str_sqmetres_eng34)
score_str_sqmetres_eng36 = re.sub ('Лист', 'Sheet', score_str_sqmetres_eng35)




mists_eng = score_str_sqmetres_eng22.split('\n')

def check_eng(score_str_kg):
	vivod = ""
	if re.search(f'{k}', score_str_kg):
		vivod += str("O-Ring")
	if re.search(f'{m}', score_str_kg):
		vivod +=str(" Cuff")
	if re.search(f'{b}' or '{g}' , score_str_kg):
		vivod +=str(' Wiper')
	if re.search(f'{g}', score_str_kg):
		vivod +=str(' Wiper')
	if re.search(f'{c}', score_str_kg):
		vivod +=str(' Cord')
	return vivod

product_name_eng = check_eng(score_str_kg)

document_eng = Document()
document_eng.add_picture('shapka_eng.jpg', width=Inches(7.086))
sertificat_eng = document_eng.add_paragraph('                                                             ')
sertificat_eng.add_run('CERTIFICATE OF QUIALITY №').bold = True
sertificat_eng.add_run(str(date_today)).bold = True
name_of_product_eng = document_eng.add_paragraph ("")
name_of_product_eng.add_run ("Name of production: ").bold = True
name_of_product_eng.add_run (product_name_eng).underline = True
partiya_eng = document_eng.add_paragraph('')
partiya_eng.add_run ("Batch № " + partiya_date).bold = True

passport_table_eng = document_eng.add_table (rows = len_of_future_tables+2, cols = 5)
passport_table_eng.style = 'Table Grid'
set_column_width(passport_table_eng, 0, 0.27)
set_column_width(passport_table_eng, 1, 4.5)
set_column_width(passport_table_eng, 2, 2)
set_column_width(passport_table_eng, 3, 1.5)
set_column_width(passport_table_eng, 4, 0.3)
hdr_cells_eng = passport_table_eng.rows[0].cells
hdr_cells_eng[0].text = '№'
hdr_cells_eng[1].text = 'Name'
hdr_cells_eng[2].text = 'Material'
hdr_cells_eng[3].text = 'Batch №'
hdr_cells_eng[4].text = 'Pcs'

for row in range(1,len_of_future_tables+2):
    for col in range(1):
    	cell = passport_table_eng.cell(row, col)
    cell.text = str(row)

for row in range(1,len_of_future_tables+2):
	if row == 10:
		break
	for col in range(4):
    		cell = passport_table_eng.cell(row, col)
	cell.text = str(partiya_date) + '-00' + str(row)
if len_of_future_tables > 9:
	for row in range(10,len_of_future_tables+2):
		for col in range(4):
				cell = passport_table_eng.cell(row, col)
		cell.text = str(partiya_date) + '-0' + str(row)


for row, mist in zip(range(1, len_of_future_tables+2), mists_eng):
    for col in range(2):
        cell = passport_table_eng.cell(row, col)
    cell.text = (mist)

for row, count in zip(range(1, len_of_future_tables+2), counts_kg_new_eng):
    for col in range(5):
        cell = passport_table_eng.cell(row, col)
    cell.text = (count)
for row, mater_eng in zip(range(1, len_of_future_tables+2), material_eng):
    for col in range(3):
        cell = passport_table_eng.cell(row, col)
    cell.text = (mater_eng)

rezultat_of_control_eng = document_eng.add_paragraph('\n                                                             ')
rezultat_of_control_eng.add_run('1.  Result of control').bold = True

fizikal_mech_char2_eng = document_eng.add_paragraph('')
fizikal_mech_char2_eng.add_run('1.1 The physical and mechanical properties of the material correspond: \n' + fizikal_mech_gosts_eng + "\nDimensions and appearance meet the requirements of" + " " + fizikal_mech_gosts_eng).bold = True
enough12 = document_eng.add_paragraph('\n                                              ')
enough12.add_run('Conclusion of the technical control department').bold = True
enough22 = document_eng.add_paragraph('\n        ')
enough22.add_run(' Batch № ' + str(partiya_date) + " meets the requirements ").underline = True
enough22.add_run(fizikal_mech_gosts_eng).bold = True
months2_eng_list = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
mesyac2_eng = months2_eng_list[(int(month_today)-1) % 12]
enough32 = document_eng.add_paragraph ("")
enough32.add_run ('\nThe warranty period starts from ' + mesyac2_eng + " 20" + datetime.datetime.today().strftime("%y") + "\nShelf life - 5 years").italic = True
enough42 = document_eng.add_paragraph ('\nHead of control                                                                                                   Kuznetsov Alexey')

document_eng.add_page_break()
document_eng.save('Паспорт на английском для счета номер ' + number_of_score.group(0) + '.docx')

print ("Паспорт на английском готов")





print ("Нажмите любую клавишу для завершения")
input ('')